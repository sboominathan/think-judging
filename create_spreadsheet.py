import csv
import json
from docx import Document

def get_submitted_apps():
	''' 
	Get list of all JSON records corresponding to submitting applications
	'''

	APP_FILE_NAME = "think_apps_1_1.json"

	 # Read in applications JSON file
	with open(APP_FILE_NAME) as json_file:  
		data = []
		for line in json_file:
			data.append(json.loads(line))

	# Get list of all applications with non-empty file name submission
	submitted_apps = []
	for application in data:
		if len(application["projectSubmission"]["fileName"]) > 0:
			submitted_apps.append(application)

	return submitted_apps

def create_judging_spreadsheet():
	'''
	Convert JSON file into CSV spreadsheet for judging
	'''

	submitted_apps = get_submitted_apps()
	print("# of Applications: ", len(submitted_apps))

	OUTPUT_FILE_NAME = "apps_cleaned_1_1.csv"
	bad_submissions = 0

	# Write username, project title, project areas, and Dropbox link to CSV file
	with open(OUTPUT_FILE_NAME, 'w') as csvfile:
		application_writer = csv.writer(csvfile, delimiter=',')
		application_writer.writerow(["Username", "Title", "Project Areas", "Dropbox Link", "Submitted?", "Partner?"])
		for app in submitted_apps:
			username = app["username"]
			project_title = app["projectSubmission"]["projectTitle"]
			areas = app["projectQuestions"]["projectAreas"]
			project_link = app["projectSubmission"]["fileLink"]
			submitted = app["projectSubmission"]["submitted"]
			partner = app["projectSubmission"]["partner"]

			application_writer.writerow([username, project_title, areas, project_link, submitted, partner])

			if len(project_link) == 0:
				print("Bad Username: {}".format(username))
				bad_submissions += 1

	print("# of Bad Submissions: ", bad_submissions)
			
def create_response_doc():
	'''
	Create Word doc containing all student responses to 3 questions in application, along 
	with additional info
	'''
	submitted_apps = get_submitted_apps()
	responses_doc = Document()

	for app in submitted_apps:
		write_user_responses(responses_doc, app)
		responses_doc.add_page_break()

	responses_doc.save("student_responses.docx")
		

def write_user_responses(document, application):
	'''
	Adds responses from a particular applicant into the provided Word document
	'''
	username = application["username"]
	project_title = application["projectSubmission"]["projectTitle"]

	document.add_heading(username + ": " + project_title, level=3)
	document.add_heading("What kind of technical skills will be needed to complete your project?", level=4)
	document.add_paragraph(application["projectQuestions"]["skills"])

	document.add_heading("What do you anticipate being the biggest obstacles in your project?", level=4)
	document.add_paragraph(application["projectQuestions"]["obstacles"])

	document.add_heading("Where do you believe that THINK mentorship can benefit your project?", level=4)
	document.add_paragraph(application["projectQuestions"]["mentorship"])

	document.add_heading("Additional Info", level=4)
	document.add_paragraph(application["projectSubmission"]["additional"])

if __name__ == '__main__':
	# create_judging_spreadsheet()
	create_response_doc()
