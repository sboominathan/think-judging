import csv
from docx import Document
# from testPlotLib import graph_results

def parse_results(print_bin=None):
    TOTAL_PAPER_COUNT = 277
    with open("judging_results_initial_2019.csv", "rt") as results_file:
         parsed_results = {}
         raw_results = csv.reader(results_file)
         next(raw_results, None)

         judge_counts = {}
         paper_count = {}
         judge_responses = {}

         for row in raw_results:
             insert_results(parsed_results, row, judge_counts, paper_count, judge_responses)

         # print("# of Papers Read: {}".format(len(parsed_results)))
         one_judge_papers = [paper for paper in parsed_results.keys() if paper_count[paper] == 1]
         two_judge_papers = [paper for paper in parsed_results.keys() if paper_count[paper] == 2]

         # print("# One judge papers: {}".format(len(one_judge_papers)))
         # print("# Two judge papers: {}".format(len(two_judge_papers)))
         # print("# Zero judge papers: {}".format(TOTAL_PAPER_COUNT-(len(one_judge_papers) + len(two_judge_papers))))
         # print("---------------------")

         # Debugging which papers don't have identical titles
         # for paper, count in paper_count.items():
         #    if count == 1:
         #        print(paper)

         print('Total Number of Submissions: {}'.format(sum(judge_counts.values())))
         bins = separate_results(parsed_results)

         # for bin in sorted(bins.keys()):
         #    print("Bin: {}, ".format(bin) + "# papers: {}".format(len(bins[bin])))
         if print_bin:
             print("Papers in {} Bin".format(print_bin))
             for paper, results in bins[print_bin].items():
                print('Title: {}'.format(paper))
                # print('Judge 1: {} - {}'.format(results['judge1']['judge'], results['judge1']['decision']))
                # print('Judge 2: {} - {}'.format(results['judge2']['judge'], results['judge2']['decision']))
            
                # print('-------------------')
             print('===========================================================================')
             
         # for bin in bins:
         #     write_output_new(bins, bin)



def insert_results(results, row, judge_counts, paper_count, judge_responses):
    title = row[2].lower()
    judge_number = "judge1" if title not in results else "judge2"

    if title not in results:
        results[title] = {}
        paper_count[title] = 0
    paper_count[title] += 1

    judge_first = row[1].split(" ")[0]

    if judge_first == 'Tiffany':
        judge_first = row[1]
    elif judge_first == 'Mau':
        judge_first = 'Mauri'
        

    if judge_first not in judge_counts:
        judge_counts[judge_first] = 0
        judge_responses[judge_first] = {'Yes': 0, 'Maybe': 0, 'No': 0}

    judge_counts[judge_first] += 1
    judge_responses[judge_first][row[13]] += 1

    # results[title][judge_number] = {"judge": judge_first, "area": row[3], "summary": row[4], "impact": row[5], \
    #                                  "innovation": row[6], "clarity": row[7], "feasibility": row[8], \
    #                                 "gen-notes": row[9], "benefit": row[10], "good": row[11], "bad": row[12],\
    #                                  "decision": row[13], "addl-notes": row[14]}
    assert(row[13] in ["Yes","No","Maybe"])
    results[title][judge_number] = {"judge": judge_first, "decision": row[13]}

def separate_results(parsed_results):
    bins = {}
    for paper in parsed_results:
        if "judge2" not in parsed_results[paper]:
            decision = parsed_results[paper]["judge1"]["decision"]
            decision_map = {"Yes": '1y', "Maybe": '1m', "No": '1n'}
            bins.setdefault(decision_map[decision],{})[paper] = parsed_results[paper]
            continue

        decisions = sorted([parsed_results[paper]["judge1"]["decision"], parsed_results[paper]["judge2"]["decision"]])
        if decisions == ["Yes","Yes"]:
            bins.setdefault('2y',{})[paper] = parsed_results[paper]
        elif decisions == ["Maybe","Yes"]:
            bins.setdefault('1y1m',{})[paper] = parsed_results[paper]
        elif decisions == ["No","Yes"]:
            bins.setdefault('1n1y',{})[paper] = parsed_results[paper]
        elif decisions == ["Maybe","Maybe"]:
            bins.setdefault('2m',{})[paper] = parsed_results[paper]
        elif decisions == ["Maybe","No"]:
            bins.setdefault('1m1n',{})[paper] = parsed_results[paper]
        elif decisions == ["No","No"]:
            bins.setdefault('2n',{})[paper] = parsed_results[paper]
        else:
            print("No bins fit: {}".format(paper))
    return bins

def write_output_new(bins, bin_name):
      paper_list = bins[bin_name]
      document = Document()
      # graph_results()
      for paper in paper_list:
         document.add_heading(paper.title(), 0)
         document.add_heading('Judge 1: ' + paper_list[paper]["judge1"]["judge"] , level=1)
         write_judging_results_new(document, paper_list[paper]["judge1"])
         document.add_heading('Judge 2: ' + paper_list[paper]["judge2"]["judge"], level=1)
         write_judging_results_new(document, paper_list[paper]["judge2"])
         # document.add_picture(paper+'.png')
         document.add_page_break()
      document.save(bin_name + "_2018.docx")

def write_judging_results_new(document, results):
    document.add_heading("Area: " + results["area"], level=2)

    document.add_heading("Summary: ", level=2)
    document.add_paragraph(results["summary"])

    table = document.add_table(rows=1, cols=4, style= "Light Shading Accent 1")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Impact'
    hdr_cells[1].text = 'Innovation'
    hdr_cells[2].text = 'Clarity'
    hdr_cells[3].text = 'Feasibility'
    row_cells = table.add_row().cells
    row_cells[0].text = results["impact"]
    row_cells[1].text = results["innovation"]
    row_cells[2].text = results["clarity"]
    row_cells[3].text = results["feasibility"] 
            
    document.add_heading("General Notes: ", level=2)
    document.add_paragraph(results["gen-notes"])

    document.add_heading("Personal Benefit to Applicant: ", level=2)
    document.add_paragraph(results["benefit"])

    document.add_heading("What's Good: ", level=2)
    document.add_paragraph(results["good"])

    document.add_heading("What's Bad: ", level=2)
    document.add_paragraph(results["bad"])

    document.add_heading("Decision: " + results["decision"], level=1)

    document.add_heading("Final Notes: ", level=2)
    document.add_paragraph(results["addl-notes"])


if __name__ == "__main__":
    parse_results(print_bin='2y')
    parse_results(print_bin='1y1m')
    parse_results(print_bin='2m')
    parse_results(print_bin='1n1y')
