import pandas as pd
from collections import defaultdict

from tqdm import tqdm
from docx import Document

RESPONSES_PATH = "PUT CSV PATH HERE"

def get_buckets_judging_responses(responses_csv_path):

    '''
        Groups papers by judge responses.

        Returns:
         - Pandas DataFrame with all judging responses
         - Dict mapping each response category (i.e, 2 no, 2 yes, etc.) to paper titles in that bucket
    '''
    
    # Ignore first row with dummy responses
    responses_df = pd.read_csv(responses_csv_path).iloc[1:]
    
    # Rename columns for brevity
    responses_df.columns = ['timestamp', 'judge', 'title', 'summary', 'impact', 'innovation',
                             'clarity', 'feasibility', 'benefit', 'good', 'bad', 'rd2', 'notes']
    
    response_to_int_map = {'Yes': 1, 'Maybe': 0, 'No': -1}
    one_judge_map = {(1,): '1y', (0,): '1m', (-1,): '1n'}
    two_judge_map = {(1,1): '2y', (-1,-1): '2n', (0,0): '2m',
                    (-1,1): '1y1n', (0,1): '1y1m', (-1,0): '1m1n'}
    
    responses_df['rd2_final'] = responses_df.apply(lambda x: response_to_int_map[x['rd2']], axis=1)
    
    # Create title buckets
    response_buckets = defaultdict(list)
    
    # Group responses by title
    responses_grouped_title = responses_df.groupby('title')
    for title, responses in tqdm(responses_grouped_title):
        total_response = tuple(sorted(responses['rd2_final'].values))
        if len(total_response) == 1:
            bucket = one_judge_map[total_response]
        elif len(total_response) == 2:
            bucket = two_judge_map[total_response]
        else:
            print(f'{title} has > 2 judges')
        
        response_buckets[bucket].append(title)
    
    return responses_df, response_buckets


def write_output_new(responses_df,
                     bucket_name,
                     paper_list):

    '''
        Creates doc with all judging responses for provided list of paper titles
    '''

    document = Document()
    
    for title in paper_list:
        document.add_heading(title, 0)
        
        responses_for_title = responses_df[responses_df.title == title]
        
        judge1_response = responses_for_title.iloc[0]
        document.add_heading(f"Judge 1: {judge1_response['judge']}",
                             level=1)
        write_judging_results_new(document, judge1_response)
        
        if len(responses_for_title) > 1:
            judge2_response = responses_for_title.iloc[1]
            document.add_heading(f"Judge 2: {judge2_response['judge']}",
                                 level=1)
            write_judging_results_new(document, judge2_response)
        
        document.add_page_break()
        
    document.save(f"{bucket_name}_2020.docx")


def write_judging_results_new(document, results):
    '''
        Writes single judge's response to document
    '''
    document.add_heading("Summary: ", level=2)
    document.add_paragraph(results["summary"])

    table = document.add_table(rows=1, cols=4, style= "Light Shading Accent 1")

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Impact'
    hdr_cells[1].text = 'Innovation'
    hdr_cells[2].text = 'Clarity'
    hdr_cells[3].text = 'Feasibility'
    
    row_cells = table.add_row().cells
    row_cells[0].text = str(results["impact"])
    row_cells[1].text = str(results["innovation"])
    row_cells[2].text = str(results["clarity"])
    row_cells[3].text = str(results["feasibility"])
    
    document.add_heading("Personal Benefit to Applicant: ", level=2)
    document.add_paragraph(str(results["benefit"]))

    document.add_heading("What's Good: ", level=2)
    document.add_paragraph(str(results["good"]))

    document.add_heading("What's Bad: ", level=2)
    document.add_paragraph(str(results["bad"]))

    document.add_heading(f"Decision: {results['rd2']}", level=1)

    document.add_heading("Final Notes: ", level=2)
    document.add_paragraph(str(results["notes"]))


def generate_response_docs(responses_path_csv):
    '''
        Generates docs of responses for each bucket of possible responses
    '''
    responses_df, response_buckets = get_buckets_judging_responses(responses_path_csv)
    
    for bucket_name, paper_list in response_buckets.items():
        write_output_new(responses_df,
                         bucket_name,
                         paper_list)
        

if __name__ == '__main__':
    generate_response_docs(RESPONSES_PATH)

